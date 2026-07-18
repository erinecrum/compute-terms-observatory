"""Generic fetcher: retrieve one public document and normalize it.

One function fetches any registry `Document`. There are no per-provider classes:
providers differ only by URL, and the normalizer is deliberately generic so a
new provider needs no code.

Normalization mirrors the radar project: isolate the main content, strip
non-content noise (scripts, nav, footers) that would create phantom diffs, and
render the text one logical block per line so a single edited clause shows up as
a single changed line downstream.

We keep BOTH the raw HTML (the archival corpus — sacred, never discarded) and the
normalized text (the change signal we diff and the input we extract from).
"""

from __future__ import annotations

import io
import re
import time
from urllib import robotparser
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup

from .model import Document, FetchResult, sha256_bytes, sha256_text
from .textcheck import looks_like_text

# We identify honestly as an archival bot (not a spoofed browser) and link back to
# the project, so any site owner can see exactly who we are. We are archivists,
# not scrapers evading anyone.
_UA = (
    "ComputeTermsObservatory/1.0 "
    "(+https://github.com/erinecrum/compute-terms-observatory; public terms archival)"
)
_HEADERS = {
    "User-Agent": _UA,
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
}
_TIMEOUT = 45
_RETRIES = 3  # polite exponential back-off on transient network/5xx errors

# robots.txt parsers, cached per host so we read each site's rules once per run.
_robots_cache: dict = {}


def _robots_allows(url: str) -> tuple:
    """Return (allowed, reason). We honor robots.txt: if a host disallows a path
    for us, we do not fetch it and record the refusal as a failure. A missing or
    unreadable robots.txt is treated as 'allowed' (the web default)."""
    parsed = urlparse(url)
    host = f"{parsed.scheme}://{parsed.netloc}"
    if host not in _robots_cache:
        rp = None
        try:
            r = requests.get(urljoin(host, "/robots.txt"), headers=_HEADERS, timeout=20)
            if r.status_code == 200:
                rp = robotparser.RobotFileParser()
                rp.parse(r.text.splitlines())
        except requests.RequestException:
            rp = None  # unreadable robots.txt -> treat as allowed
        _robots_cache[host] = rp
    rp = _robots_cache[host]
    if rp is None:
        return True, ""
    return (rp.can_fetch(_UA, url), "disallowed by robots.txt")


def _get_politely(url: str) -> requests.Response:
    """One GET with exponential back-off on transient errors. Callers fetch
    sequentially (one request at a time) so we never hammer a host in parallel."""
    last = None
    for attempt in range(_RETRIES):
        try:
            resp = requests.get(url, headers=_HEADERS, timeout=_TIMEOUT)
            if resp.status_code >= 500:
                resp.raise_for_status()
            return resp
        except requests.RequestException as exc:
            last = exc
            if attempt < _RETRIES - 1:
                time.sleep(2 ** attempt)  # 1s, 2s
    raise last

# Preferred content containers, tried in order, with a whole-body fallback so a
# template tweak never blanks a snapshot.
_CONTENT_SELECTORS = [
    "main",
    "article",
    "[role=main]",
    "#main-content",
    "#content",
    ".content",
    "#aws-page-content-main",
]

# Elements that are never legal content but often sit inside the main container.
_NOISE = "script, style, noscript, svg, nav, header, footer, form, iframe"

# Below this many characters we assume we were blocked or the layout broke, and
# refuse to record the snapshot — a near-empty snapshot would read as "the
# provider deleted its entire terms document" on the next diff.
_MIN_CHARS = 200


def _lineify(raw_text: str) -> str:
    """Squeeze whitespace and drop blank lines so a single edited clause shows up
    as a single changed line downstream. Shared by the HTML and PDF paths."""
    lines = []
    for line in raw_text.split("\n"):
        line = re.sub(r"\s+", " ", line).strip()
        if line:
            lines.append(line)
    return "\n".join(lines)


def _pdf_text(content: bytes) -> str:
    """Extract text from a PDF document, rendered line-oriented like the HTML path.
    Used for documents published only as PDFs (e.g. the Microsoft Customer
    Agreement), so they diff and extract consistently with HTML sources."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return _lineify("\n".join(pages))


def _docx_text(content: bytes) -> str:
    """Extract text from a DOCX document, including table cells — SLA credit tiers
    are usually laid out in tables, so paragraph text alone would miss them. Used
    for documents published only as DOCX (e.g. the Azure consolidated SLA)."""
    from docx import Document as DocxDocument

    d = DocxDocument(io.BytesIO(content))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return _lineify("\n".join(parts))


def _normalize(html: str) -> str:
    """Extract the document text and render it line-oriented for clean diffing."""
    soup = BeautifulSoup(html, "html.parser")

    container = None
    for sel in _CONTENT_SELECTORS:
        container = soup.select_one(sel)
        if container is not None:
            break
    if container is None:
        container = soup.body or soup  # last resort: whole document

    for tag in container.select(_NOISE):
        tag.decompose()

    return _lineify(container.get_text("\n"))


# Browser-tier UA: a real headless-Chromium request (so JS-challenge pages render)
# that STILL discloses who we are by appending our bot identifier. We present as a
# browser because we drive one, but we never hide that we are an archival bot.
_BROWSER_UA = (
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/125.0.0.0 Safari/537.36 " + _UA
)

# Markers that indicate an interactive challenge / CAPTCHA. Hard rule: we never
# attempt to solve one — we detect it, log it, and fall back to the wayback tier.
_CHALLENGE_MARKERS = (
    "captcha",
    "cf-challenge",
    "just a moment",
    "verify you are human",
    "attention required",
    "checking your browser",
    "enable javascript and cookies to continue",
)

_WAYBACK_AVAILABLE = "https://archive.org/wayback/available"


def _finalize(result: FetchResult, url: str, content: bytes, content_type: str,
              html_text: str | None = None, charset: str = "") -> None:
    """Detect the document type, normalize to text, and populate the result. Shared
    by all fetch tiers so a browser- or wayback-sourced document goes through the
    exact same selector/filter/version pipeline as a direct fetch. Raises on an
    empty/garbage capture so we never store a misleading snapshot."""
    content_type = (content_type or "").lower()
    is_pdf = content[:5] == b"%PDF-" or "application/pdf" in content_type or url.lower().endswith(".pdf")
    is_docx = url.lower().endswith(".docx") or "wordprocessingml" in content_type
    is_other_office = (content[:4] == b"PK\x03\x04" and not is_docx) or url.lower().endswith((".xlsx", ".pptx"))

    if is_pdf:
        text = _pdf_text(content)
        result.raw_bytes, result.raw_ext, result.raw_sha256 = content, "pdf", sha256_bytes(content)
    elif is_docx:
        text = _docx_text(content)
        result.raw_bytes, result.raw_ext, result.raw_sha256 = content, "docx", sha256_bytes(content)
    elif is_other_office:
        raise RuntimeError("unsupported document format (Office binary other than DOCX/PDF)")
    else:
        raw_html = html_text if html_text is not None else content.decode("utf-8", "replace")
        # Refuse to snapshot a capture that did not decode as text (binary,
        # compressed, or wrong charset). Record the content-type and encoding so
        # the underlying fetch problem can be diagnosed and fixed, not just
        # suppressed downstream.
        is_text, stats = looks_like_text(raw_html)
        if not is_text:
            enc = charset or "unknown"
            result.meta["reject_reason"] = "non_text"
            result.meta["content_type"] = content_type or "unknown"
            result.meta["encoding"] = enc
            result.meta["replacement_ratio"] = stats["replacement_ratio"]
            result.meta["nonprint_ratio"] = stats["nonprint_ratio"]
            raise RuntimeError(
                "response decoded as non-text (likely binary, compressed, or wrong "
                f"charset); content-type={content_type or 'unknown'} encoding={enc} "
                f"replacement_ratio={stats['replacement_ratio']} "
                f"nonprint_ratio={stats['nonprint_ratio']} — refusing to snapshot"
            )
        text = _normalize(raw_html)
        result.raw_html, result.raw_ext, result.raw_sha256 = raw_html, "html", sha256_text(raw_html)

    if len(text) < _MIN_CHARS:
        raise RuntimeError(
            f"content suspiciously short ({len(text)} chars) — likely blocked or the "
            "page layout changed; refusing to snapshot"
        )
    result.text = text
    result.text_sha256 = sha256_text(text)
    result.char_count = len(text)
    result.ok = True


def _new_result(doc: Document) -> FetchResult:
    return FetchResult(
        provider=doc.provider, provider_name=doc.provider_name, doc_type=doc.doc_type,
        name=doc.name, url=doc.url, slug=doc.slug, fetch_method=doc.fetch_method,
    )


def _fetch_direct(doc: Document, result: FetchResult) -> FetchResult:
    """Tier 1: a plain, polite HTTP GET as our honest archival bot."""
    try:
        allowed, reason = _robots_allows(doc.url)
        if not allowed:
            result.ok, result.error = False, reason
            return result
        resp = _get_politely(doc.url)
        result.http_status = resp.status_code
        resp.raise_for_status()
        result.fetch_method = "direct"
        _finalize(result, doc.url, resp.content, resp.headers.get("content-type", ""),
                  resp.text, charset=resp.encoding or "")
    except Exception as exc:  # noqa: BLE001
        result.ok, result.error = False, f"{type(exc).__name__}: {exc}"
    return result


def _fetch_browser(doc: Document, result: FetchResult) -> FetchResult:
    """Tier 2: drive headless Chromium (for JS-rendered / challenge pages), then run
    the rendered HTML through the same pipeline. Never solves a CAPTCHA."""
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        result.ok = False
        result.error = "playwright not installed (browser tier available only in the workflow)"
        return result
    try:
        allowed, reason = _robots_allows(doc.url)
        if not allowed:
            result.ok, result.error = False, reason
            return result
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            ctx = browser.new_context(user_agent=_BROWSER_UA)
            page = ctx.new_page()
            resp = page.goto(doc.url, wait_until="networkidle", timeout=_TIMEOUT * 1000)
            page.wait_for_timeout(1500)  # let late content settle
            html = page.content()
            result.http_status = resp.status if resp else None
            browser.close()
        low = html.lower()
        if any(m in low for m in _CHALLENGE_MARKERS):
            # HARD RULE: do not attempt to bypass/solve. Log and fall back to wayback.
            result.ok = False
            result.error = "interactive challenge/CAPTCHA detected; not bypassing (hard rule)"
            result.meta["challenge"] = True
            return result
        result.fetch_method = "browser"
        _finalize(result, doc.url, html.encode("utf-8"), "text/html",
                  html_text=html, charset="utf-8")
    except Exception as exc:  # noqa: BLE001
        result.ok, result.error = False, f"browser: {type(exc).__name__}: {exc}"
    return result


def _save_page_now(url: str) -> None:
    """Ask the Internet Archive to capture the URL now, so future runs have a fresh
    capture to fall back to. Fire-and-forget: we do not block on or fail over it."""
    try:
        requests.get(f"https://web.archive.org/save/{url}", headers=_HEADERS, timeout=8)
    except requests.RequestException:
        pass


def _fetch_wayback(doc: Document, result: FetchResult) -> FetchResult:
    """Tier 3: fall back to the most recent Internet Archive capture. The version is
    dated by CAPTURE time (recorded in metadata), not our fetch time. Also triggers
    a fresh Save Page Now so future captures stay current."""
    try:
        avail = requests.get(_WAYBACK_AVAILABLE, params={"url": doc.url},
                             headers=_HEADERS, timeout=30).json()
        snap = (avail.get("archived_snapshots") or {}).get("closest") or {}
        if not snap.get("available"):
            result.ok, result.error = False, "wayback: no capture available"
            _save_page_now(doc.url)
            return result
        # `id_` returns the ORIGINAL captured bytes without the Wayback UI/toolbar.
        raw_url = re.sub(r"/web/(\d+)/", r"/web/\1id_/", snap["url"], count=1)
        resp = _get_politely(raw_url)
        result.http_status = resp.status_code
        result.fetch_method = "wayback"
        _finalize(result, doc.url, resp.content, resp.headers.get("content-type", ""),
                  resp.text, charset=resp.encoding or "")
        ts = snap.get("timestamp", "")
        if len(ts) == 14:
            result.capture_timestamp = f"{ts[0:4]}-{ts[4:6]}-{ts[6:8]}T{ts[8:10]}:{ts[10:12]}:{ts[12:14]}Z"
        result.meta["wayback_capture"] = result.capture_timestamp
        result.meta["wayback_url"] = snap["url"]
    except Exception as exc:  # noqa: BLE001
        result.ok, result.error = False, f"wayback: {type(exc).__name__}: {exc}"
    _save_page_now(doc.url)  # keep future captures fresh regardless
    return result


def fetch_document(doc: Document) -> FetchResult:
    """Fetch one document via its configured tier, with fallback. Returns a
    FetchResult with ok=False (never raises) so one broken URL can't abort a run.

    Tiers: direct (default). browser (headless Chromium) for pages that block or
    require JS; on browser failure OR a detected CAPTCHA it falls back to wayback.
    wayback goes straight to the Internet Archive capture.
    """
    result = _new_result(doc)
    method = doc.fetch_method

    if method == "direct":
        return _fetch_direct(doc, result)

    if method == "browser":
        r = _fetch_browser(doc, result)
        if r.ok:
            return r
        # Browser failed or hit a challenge -> wayback fallback (points 3 & 4).
        wb = _fetch_wayback(doc, _new_result(doc))
        if wb.ok:
            wb.meta["fell_back_from"] = "browser"
            wb.meta["browser_error"] = r.error
            return wb
        r.error = f"browser failed ({r.error}); wayback fallback also failed ({wb.error})"
        return r

    if method == "wayback":
        return _fetch_wayback(doc, result)

    result.ok, result.error = False, f"unknown fetch_method {method!r}"
    return result
